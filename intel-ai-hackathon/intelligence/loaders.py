from __future__ import annotations

import email
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterator

from docx import Document


@dataclass
class LoadedDocument:
    doc_id: str
    source_type: str  # report | email | whatsapp
    title: str
    source_file: str
    text: str
    occurred_at: datetime | None


_WHATSAPP_LINE = re.compile(
    r"^\s*(\d{1,2}[/.-]\d{1,2}[/.-]\d{2,4})[,\s]+(\d{1,2}:\d{2}(?::\d{2})?(?:\s*[AP]M)?)\s*[-–—]\s*([^:]+):\s*(.*)$",
    re.IGNORECASE,
)


def _parse_whatsapp_dt(date_part: str, time_part: str) -> datetime | None:
    for fmt in ("%d/%m/%Y %H:%M", "%d-%m-%Y %H:%M", "%d.%m.%Y %H:%M", "%m/%d/%Y %H:%M"):
        try:
            return datetime.strptime(f"{date_part} {time_part.split()[0][:5]}", fmt)
        except ValueError:
            continue
    return None


def load_whatsapp(path: Path) -> LoadedDocument:
    raw = path.read_text(encoding="utf-8", errors="replace")
    lines = raw.splitlines()
    first_dt: datetime | None = None
    normalized: list[str] = []
    for line in lines:
        m = _WHATSAPP_LINE.match(line)
        if m:
            d, t, who, msg = m.groups()
            dt = _parse_whatsapp_dt(d, t)
            if dt and first_dt is None:
                first_dt = dt
            normalized.append(f"[{d} {t}] {who.strip()}: {msg.strip()}")
        else:
            normalized.append(line)
    text = "\n".join(normalized)
    title = path.stem
    return LoadedDocument(
        doc_id=f"whatsapp:{path.name}",
        source_type="whatsapp",
        title=title,
        source_file=path.name,
        text=text,
        occurred_at=first_dt,
    )


def load_email(path: Path) -> LoadedDocument:
    raw_bytes = path.read_bytes()
    occurred_at: datetime | None = None
    text: str
    title: str
    try:
        msg = email.message_from_bytes(raw_bytes)
        subj = msg.get("Subject", path.stem) or path.stem
        date_hdr = msg.get("Date")
        if date_hdr:
            try:
                occurred_at = email.utils.parsedate_to_datetime(date_hdr)
            except (TypeError, ValueError, OverflowError):
                occurred_at = None
        body_parts: list[str] = []
        if msg.is_multipart():
            for part in msg.walk():
                ctype = part.get_content_type()
                if ctype == "text/plain":
                    payload = part.get_payload(decode=True)
                    if isinstance(payload, bytes):
                        body_parts.append(payload.decode("utf-8", errors="replace"))
        else:
            payload = msg.get_payload(decode=True)
            if isinstance(payload, bytes):
                body_parts.append(payload.decode("utf-8", errors="replace"))
            elif isinstance(payload, str):
                body_parts.append(payload)
        text = "\n".join(body_parts) if body_parts else ""
        title = str(subj)
        if len(text.strip()) < 20 and msg.keys():
            text = raw_bytes.decode("utf-8", errors="replace")
    except Exception:
        text = raw_bytes.decode("utf-8", errors="replace")
        title = path.stem
    return LoadedDocument(
        doc_id=f"email:{path.name}",
        source_type="email",
        title=title,
        source_file=path.name,
        text=text,
        occurred_at=occurred_at,
    )


def load_report(path: Path) -> LoadedDocument:
    if path.suffix.lower() == ".docx":
        doc = Document(path)
        text = "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    else:
        text = path.read_text(encoding="utf-8", errors="replace")
    mtime = datetime.fromtimestamp(path.stat().st_mtime)
    return LoadedDocument(
        doc_id=f"report:{path.name}",
        source_type="report",
        title=path.name,
        source_file=path.name,
        text=text,
        occurred_at=mtime,
    )


def load_document(path: Path) -> LoadedDocument | None:
    suffix = path.suffix.lower()
    parent = path.parent.name.lower()
    if parent == "whatsapp" or "whatsapp" in path.name.lower():
        if suffix in {".txt"}:
            return load_whatsapp(path)
    if parent == "emails" or parent == "email" or suffix == ".eml":
        if suffix in {".eml", ".txt", ".msg"}:
            return load_email(path)
    if parent == "reports" or suffix in {".txt", ".md", ".rst", ".docx"}:
        if suffix in {".txt", ".md", ".rst", ".docx"}:
            return load_report(path)
    if suffix == ".eml":
        return load_email(path)
    if suffix in {".txt", ".md", ".docx"}:
        if "whatsapp" in str(path).lower():
            return load_whatsapp(path)
        return load_report(path)
    return None


def iter_documents(data_root: Path) -> Iterator[LoadedDocument]:
    for path in sorted(data_root.rglob("*")):
        if not path.is_file():
            continue
        doc = load_document(path)
        if doc is not None:
            yield doc


def default_data_root() -> Path:
    return Path(__file__).resolve().parent.parent / "data"
